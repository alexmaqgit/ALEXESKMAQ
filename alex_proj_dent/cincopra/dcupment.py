# إنشاء ملف Word فاخر لقائمة أسعار عيادة أسنان باللغة العربية
'''from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor

# إنشاء المستند
doc = Document()

# دالة لجعل الفقرة من اليمين لليسار RTL
def set_rtl(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

# عنوان العيادة
title = doc.add_paragraph("عيادة الابتسامة الحديثة لطب وتقويم الأسنان")
set_rtl(title)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.bold = True
run.font.size = Pt(20)

# سطر فرعي
subtitle = doc.add_paragraph("قائمة الأسعار الرسمية")
set_rtl(subtitle)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph("\n")

# إنشاء جدول الأسعار
services = [
    ("أول معاينة + خطة علاج", "5,000 ريال"),
    ("معاينة متابعة", "3,000 ريال"),
    ("تنظيف وإزالة جير", "20,000 ريال"),
    ("حشوة تجميلية متوسطة", "25,000 ريال"),
    ("علاج عصب ضرس", "100,000 ريال"),
    ("خلع عادي", "25,000 ريال"),
    ("خلع جراحي", "60,000 ريال"),
    ("تاج بورسلان", "130,000 ريال"),
    ("تاج زيركون", "200,000 ريال"),
    ("زرعة كاملة مع تاج", "1,100,000 ريال"),
    ("تقويم معدني كامل", "900,000 ريال"),
    ("تبييض الأسنان", "120,000 ريال"),
]

table = doc.add_table(rows=len(services)+1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# رأس الجدول
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "الخدمة"
hdr_cells[1].text = "السعر"

# تعبئة الجدول
for i, (service, price) in enumerate(services, start=1):
    row = table.rows[i].cells
    row[0].text = service
    row[1].text = price

# جعل كل الفقرات داخل الجدول RTL
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            set_rtl(paragraph)

# ملاحظة أسفل الجدول
doc.add_paragraph("\n")
note = doc.add_paragraph("ملاحظة: يمكن خصم رسوم المعاينة عند البدء بالعلاج. الأسعار قابلة للتعديل حسب الحالة.")
set_rtl(note)
note.runs[0].italic = True

# مساحة توقيع
doc.add_paragraph("\n")
signature = doc.add_paragraph("توقيع الإدارة: ____________________        ختم العيادة: ____________________")
set_rtl(signature)

# حفظ الملف
file_path = "/mnt/data/قائمة_أسعار_عيادة_أسنان_فاخرة.docx"
doc.save(file_path)

file_path
'''